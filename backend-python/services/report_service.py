"""统一报告生成服务 — 支持图纸评审、方案评审、训练、推理四种报告"""

import os
import uuid
from datetime import datetime
from enum import Enum

from config import settings


class ReportType(str, Enum):
    DRAWING_REVIEW = "drawing_review"
    PLAN_REVIEW = "plan_review"
    TRAINING = "training"
    INFERENCE = "inference"


class ReportService:
    """基于python-docx的报告生成器"""

    OUTPUT_DIR = getattr(settings, "report_dir", "./reports")

    def __init__(self):
        os.makedirs(self.OUTPUT_DIR, exist_ok=True)

    def generate(self, report_type: ReportType, title: str, sections: list[dict], metadata: dict | None = None) -> str:
        """生成报告并返回文件路径"""
        report_id = str(uuid.uuid4())[:8]
        filepath = os.path.join(self.OUTPUT_DIR, f"{report_id}.docx")

        if report_type == ReportType.DRAWING_REVIEW:
            doc = self._build_drawing_review(title, sections, metadata or {})
        elif report_type == ReportType.PLAN_REVIEW:
            doc = self._build_plan_review(title, sections, metadata or {})
        elif report_type == ReportType.TRAINING:
            doc = self._build_training_report(title, sections, metadata or {})
        else:
            doc = self._build_inference_report(title, sections, metadata or {})

        doc.save(filepath)
        return filepath

    def _build_drawing_review(self, title: str, sections: list[dict], meta: dict) -> "Document":
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(12)
        style.font.name = "SimSun"

        doc.add_heading("工程图纸评审报告", 0)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"图纸编号: {meta.get('drawing_id', '-')}")
        doc.add_paragraph(f"图纸页数: {meta.get('page_count', '-')}")
        doc.add_paragraph(f"问题总数: {meta.get('issue_count', len(sections))}")
        doc.add_paragraph("")

        if sections:
            doc.add_heading("一、发现的问题", level=1)
            for i, sec in enumerate(sections, 1):
                doc.add_heading(f"{i}. [{sec.get('severity', '')}] {sec.get('category', '')}", level=2)
                doc.add_paragraph(sec.get("description", ""))
                if sec.get("position"):
                    doc.add_paragraph(f"位置: {sec['position']}")

        doc.add_heading("二、综合评估", level=1)
        doc.add_paragraph(meta.get("summary", "详见附件"))

        doc.add_paragraph("\n--- AI自动生成，仅供参考 ---")
        return doc

    def _build_plan_review(self, title: str, sections: list[dict], meta: dict) -> "Document":
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(12)
        style.font.name = "SimSun"

        doc.add_heading("施工方案评审报告", 0)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"方案编号: {meta.get('document_id', '-')}")
        doc.add_paragraph(f"文档类型: {meta.get('document_type', '-')}")
        doc.add_paragraph(f"审查意见数: {meta.get('opinion_count', len(sections))}")
        doc.add_paragraph("")

        if sections:
            doc.add_heading("一、审查意见清单", level=1)
            table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
            hdr = table.rows[0].cells
            hdr[0].text = "严重程度"
            hdr[1].text = "引用规范"
            hdr[2].text = "审查意见"
            hdr[3].text = "修改建议"
            for sec in sections:
                row = table.add_row().cells
                row[0].text = sec.get("severity", "")
                row[1].text = sec.get("clause_ref", "")
                row[2].text = sec.get("comment", "")
                row[3].text = sec.get("suggestion", "")

        doc.add_heading("二、总体评价", level=1)
        doc.add_paragraph(meta.get("overall_assessment", "详见附件"))

        doc.add_paragraph("\n--- AI自动生成，仅供参考 ---")
        return doc

    def _build_training_report(self, title: str, sections: list[dict], meta: dict) -> "Document":
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(12)

        doc.add_heading("模型训练报告", 0)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"任务ID: {meta.get('job_id', '-')}")
        doc.add_paragraph(f"基础模型: {meta.get('base_model', '-')}")
        doc.add_paragraph(f"训练策略: {meta.get('strategy', '-')}")
        doc.add_paragraph(f"训练轮数: {meta.get('epochs', '-')}")

        doc.add_heading("一、性能指标", level=1)
        metrics = meta.get("metrics", {})
        table = doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
        table.rows[0].cells[0].text = "指标"
        table.rows[0].cells[1].text = "数值"
        for k, v in metrics.items():
            row = table.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)

        doc.add_heading("二、输出信息", level=1)
        doc.add_paragraph(f"模型路径: {meta.get('output_model_path', '-')}")

        doc.add_paragraph("\n--- AI自动生成，仅供参考 ---")
        return doc

    def _build_inference_report(self, title: str, sections: list[dict], meta: dict) -> "Document":
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(12)

        doc.add_heading("模型推理报告", 0)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"图片ID: {meta.get('image_id', '-')}")
        doc.add_paragraph(f"融合策略: {meta.get('fusion_strategy', '-')}")
        doc.add_paragraph(f"使用模型: {', '.join(meta.get('models_used', []))}")

        doc.add_heading("一、检测结果", level=1)
        if sections:
            table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
            hdr = table.rows[0].cells
            hdr[0].text = "#"
            hdr[1].text = "类别"
            hdr[2].text = "置信度"
            hdr[3].text = "边界框"
            for i, sec in enumerate(sections, 1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = sec.get("class_name", "")
                row[2].text = str(sec.get("confidence", ""))
                row[3].text = str(sec.get("bbox", ""))

        doc.add_heading("二、各模型对比", level=1)
        for model_name, dets in meta.get("per_model", {}).items():
            doc.add_paragraph(f"{model_name}: {len(dets)} 个检测")

        doc.add_paragraph("\n--- AI自动生成，仅供参考 ---")
        return doc
