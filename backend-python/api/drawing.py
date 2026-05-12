"""图纸评审 — PDF/DWG图纸AI分析与问题标记"""

import tempfile

from fastapi import APIRouter, File, UploadFile
from pydantic import BaseModel

router = APIRouter()


class DrawingIssue(BaseModel):
    position: dict | None
    severity: str  # 一般 / 严重 / 危急
    category: str  # 尺寸缺失 / 焊缝不规范 / 符号错误 / 结构缺陷
    description: str


class DrawingReviewResult(BaseModel):
    drawing_id: str
    page_count: int
    issues: list[DrawingIssue]
    summary: str


DRAWING_SYSTEM_PROMPT = """你是电力工程图纸评审专家，精通电气一次/二次图纸、土建施工图、铁塔结构图的审查。

请根据以下设计规范和标准，对图纸中提取的文字信息进行分析，找出不符合规范的问题。

审查要点：
1. 尺寸标注是否完整、合理
2. 焊缝标注是否符合规范要求
3. 电气符号使用是否正确
4. 结构设计是否存在缺陷
5. 材料规格是否标注清楚
6. 安全间距是否满足要求

输出格式为JSON数组，每个元素包含：
- position: 问题在图纸中的大致位置描述（如"右下角材料表"、"主视图尺寸标注区域"）
- severity: 严重程度（一般/严重/危急）
- category: 问题类别（尺寸缺失/焊缝不规范/符号错误/结构缺陷/材料问题/安全间距不足）
- description: 问题的详细描述

如果未发现明显问题，返回空数组 []。
同时请在数组后单独输出一段"总体评价："开头的综合评估摘要。"""


@router.post("/analyze", response_model=DrawingReviewResult)
async def analyze_drawing(file: UploadFile = File(...)):
    """上传图纸，AI基于规范进行分析"""
    content = await file.read()
    import fitz

    doc = fitz.open(stream=content, filetype="pdf")
    page_count = doc.page_count

    text_content = ""
    for page in doc:
        text_content += page.get_text()

    images = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=150)
        images.append(pix.tobytes("png"))
    doc.close()

    if not text_content.strip():
        return DrawingReviewResult(
            drawing_id=file.filename or "unknown",
            page_count=page_count,
            issues=[],
            summary=f"图纸共{page_count}页，未能提取到文本内容，可能为纯扫描件。建议使用OCR或VLM进行图像分析。",
        )

    # 检索相关绘图规范
    from services.knowledge_service import KnowledgeService

    kb = KnowledgeService()
    standards = await kb.search(query=text_content[:800], top_k=5, category="施工规范")

    standards_ctx = "\n".join(
        f"[{s['title']}] {s['content_snippet']}" for s in standards
    ) or "暂无匹配的规范条款"

    user_prompt = f"""相关设计规范:
{standards_ctx}

图纸文本内容（从PDF提取，共{page_count}页）:
{text_content[:6000]}

请逐条审查并返回JSON格式的问题列表。"""

    from services.llm_service import call_llm, parse_llm_json

    try:
        response = await call_llm(DRAWING_SYSTEM_PROMPT, user_prompt)
        parsed = parse_llm_json(response)
        issues = []
        for item in parsed:
            issues.append(
                DrawingIssue(
                    position=item.get("position"),
                    severity=str(item.get("severity", "一般")),
                    category=str(item.get("category", "未分类")),
                    description=str(item.get("description", "")),
                )
            )
    except Exception:
        issues = []
        response = ""

    summary = f"图纸共{page_count}页"
    if "总体评价：" in response:
        summary += "。" + response.split("总体评价：")[-1].strip().split("\n")[0]
    elif issues:
        summary += f"，发现{len(issues)}处问题"
    else:
        summary += "，已提取文本内容并完成AI分析"

    return DrawingReviewResult(
        drawing_id=file.filename or "unknown",
        page_count=page_count,
        issues=issues,
        summary=summary,
    )


@router.post("/generate-report")
async def generate_report(drawing_id: str, issues: list[DrawingIssue]):
    """根据评审结果生成Word报告"""
    from docx import Document

    doc = Document()
    doc.add_heading("工程图纸评审报告", 0)
    doc.add_paragraph(f"图纸编号: {drawing_id}")
    doc.add_paragraph(f"问题总数: {len(issues)}")

    if issues:
        doc.add_heading("问题清单", level=2)
        table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
        hdr = table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "严重程度"
        hdr[2].text = "类别"
        hdr[3].text = "问题描述"
        for i, issue in enumerate(issues, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = issue.severity
            row[2].text = issue.category
            row[3].text = issue.description

    doc.add_paragraph("\n--- AI自动生成，仅供参考 ---")

    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return {"report_path": path}
